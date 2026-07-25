"""
Document loaders for various file formats.

Supports PDF, Markdown, plain text, and DOCX files. Each loader extracts
the raw text content along with useful metadata (title, page count, word
count) and returns a :class:`Document` ready for chunking.

All file I/O is performed synchronously inside ``asyncio.to_thread`` so the
public API remains fully asynchronous without blocking the event loop.
"""

from __future__ import annotations

import asyncio
import re
from pathlib import Path
from typing import Any

from cognita.core.exceptions import DocumentLoadingError
from cognita.core.models import Document
from cognita.observability.logging import get_logger

logger = get_logger("cognita.ingestion.loaders")

# File extensions recognised by the loader.
_SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(
    {".pdf", ".md", ".markdown", ".txt", ".docx"}
)


class DocumentLoader:
    """Load documents from disk in a variety of formats.

    The loader is format-aware: it selects the appropriate parser based on
    the file extension and enriches every :class:`Document` with metadata
    such as title, page count, and word count.
    """

    def __init__(self) -> None:
        self._logger = get_logger("cognita.ingestion.loaders")

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #

    def supported_extensions(self) -> set[str]:
        """Return the set of file extensions this loader can handle."""
        return set(_SUPPORTED_EXTENSIONS)

    async def load(self, path: str | Path) -> Document:
        """Asynchronously load a single file into a :class:`Document`."""
        p = Path(path)
        return await asyncio.to_thread(self._load_sync, p)

    async def load_directory(
        self, dir_path: str | Path, pattern: str = "**/*"
    ) -> list[Document]:
        """Load every supported file beneath *dir_path* matching *pattern*."""
        p = Path(dir_path)
        return await asyncio.to_thread(self._load_directory_sync, p, pattern)

    # ------------------------------------------------------------------ #
    # Synchronous implementation (runs in a worker thread)
    # ------------------------------------------------------------------ #

    def _load_sync(self, path: Path) -> Document:
        if not path.exists():
            raise DocumentLoadingError(
                f"File not found: {path}", file_path=str(path)
            )
        if not path.is_file():
            raise DocumentLoadingError(
                f"Not a regular file: {path}", file_path=str(path)
            )

        ext = path.suffix.lower()
        if ext not in _SUPPORTED_EXTENSIONS:
            raise DocumentLoadingError(
                f"Unsupported file type '{ext}'. "
                f"Supported: {sorted(_SUPPORTED_EXTENSIONS)}",
                file_path=str(path),
            )

        try:
            if ext == ".pdf":
                content, meta = self._load_pdf(path)
            elif ext in (".md", ".markdown"):
                content, meta = self._load_markdown(path)
            elif ext == ".txt":
                content, meta = self._load_txt(path)
            elif ext == ".docx":
                content, meta = self._load_docx(path)
            else:  # pragma: no cover - guarded by the extension check above
                raise DocumentLoadingError(
                    f"Unsupported file type: {ext}", file_path=str(path)
                )
        except DocumentLoadingError:
            raise
        except Exception as exc:
            raise DocumentLoadingError(
                f"Failed to load {ext} file: {exc}", file_path=str(path)
            ) from exc

        # ``title`` is consumed separately by ``Document.from_file``; the
        # remaining metadata is attached to the Document.
        title = meta.pop("title", None)
        document = Document.from_file(
            path,
            content,
            title=title,
            metadata=meta,
        )
        self._logger.info(
            "Document loaded",
            file=str(path),
            file_type=document.file_type,
            chars=len(content),
        )
        return document

    def _load_directory_sync(
        self, path: Path, pattern: str
    ) -> list[Document]:
        if not path.exists():
            raise DocumentLoadingError(
                f"Directory not found: {path}", file_path=str(path)
            )
        if not path.is_dir():
            raise DocumentLoadingError(
                f"Not a directory: {path}", file_path=str(path)
            )

        documents: list[Document] = []
        for file_path in sorted(path.glob(pattern)):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in _SUPPORTED_EXTENSIONS:
                continue
            try:
                doc = self._load_sync(file_path)
                documents.append(doc)
            except DocumentLoadingError as exc:
                self._logger.warning(
                    "Skipping file during directory load",
                    file=str(file_path),
                    error=str(exc),
                )
        self._logger.info(
            "Directory loaded", path=str(path), documents=len(documents)
        )
        return documents

    # ------------------------------------------------------------------ #
    # Format-specific loaders  ->  (content, metadata)
    # ------------------------------------------------------------------ #

    def _load_pdf(self, path: Path) -> tuple[str, dict[str, Any]]:
        """Extract text from a PDF using ``pypdf``."""
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover
            raise DocumentLoadingError(
                "pypdf is not installed. Run: pip install pypdf",
                file_path=str(path),
            ) from exc

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        content = "\n".join(pages)

        title = self._extract_title(content)
        meta: dict[str, Any] = {
            "title": title,
            "page_count": page_count,
            "word_count": len(content.split()),
        }
        self._logger.debug(
            "PDF parsed", path=str(path), pages=page_count, chars=len(content)
        )
        return content, meta

    def _load_markdown(self, path: Path) -> tuple[str, dict[str, Any]]:
        """Convert Markdown to plain text via ``markdown`` + BeautifulSoup."""
        raw_text = self._read_text_file(path)

        # Capture the first heading as the title *before* stripping.
        title = self._extract_markdown_title(raw_text)

        try:
            import markdown as md_lib
            from bs4 import BeautifulSoup
        except ImportError as exc:  # pragma: no cover
            raise DocumentLoadingError(
                "markdown and beautifulsoup4 are required. "
                "Run: pip install markdown beautifulsoup4",
                file_path=str(path),
            ) from exc

        html = md_lib.markdown(raw_text)
        content = BeautifulSoup(html, "html.parser").get_text(separator="\n")
        content = content.strip()

        if not title:
            title = self._extract_title(content)
        meta: dict[str, Any] = {
            "title": title,
            "word_count": len(content.split()),
        }
        self._logger.debug(
            "Markdown parsed", path=str(path), chars=len(content)
        )
        return content, meta

    def _load_txt(self, path: Path) -> tuple[str, dict[str, Any]]:
        """Read a plain-text file with encoding detection and BOM stripping."""
        content = self._read_text_file(path)
        content = content.strip()
        title = self._extract_title(content)
        meta: dict[str, Any] = {
            "title": title,
            "word_count": len(content.split()),
        }
        self._logger.debug(
            "Text file parsed", path=str(path), chars=len(content)
        )
        return content, meta

    def _load_docx(self, path: Path) -> tuple[str, dict[str, Any]]:
        """Extract text from a Word document using ``python-docx``."""
        try:
            import docx
        except ImportError as exc:  # pragma: no cover
            raise DocumentLoadingError(
                "python-docx is not installed. Run: pip install python-docx",
                file_path=str(path),
            ) from exc

        doc = docx.Document(str(path))
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables (cells joined with ' | ' per row)
        for table in doc.tables:
            for row in table.rows:
                row_texts = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_texts:
                    parts.append(" | ".join(row_texts))

        content = "\n".join(parts)

        # Prefer the document's core-properties title when available.
        title = ""
        core_props = getattr(doc, "core_properties", None)
        if core_props is not None and core_props.title:
            title = core_props.title.strip()
        if not title:
            title = self._extract_title(content)

        meta: dict[str, Any] = {
            "title": title,
            "word_count": len(content.split()),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        self._logger.debug(
            "DOCX parsed", path=str(path), chars=len(content)
        )
        return content, meta

    # ------------------------------------------------------------------ #
    # Helpers
    # ------------------------------------------------------------------ #

    def _read_text_file(self, path: Path) -> str:
        """Read a text file trying UTF-8 first, then detected encoding."""
        raw = path.read_bytes()

        # Strip a leading UTF-8 BOM if present.
        if raw.startswith(b"\xef\xbb\xbf"):
            raw = raw[3:]

        # Fast path: UTF-8.
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            pass

        # Fall back to encoding detection.
        encoding = self._detect_encoding(raw)
        if encoding:
            try:
                return raw.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                pass

        # Last resort: decode as UTF-8 replacing undecodable bytes.
        return raw.decode("utf-8", errors="replace")

    @staticmethod
    def _detect_encoding(raw: bytes) -> str | None:
        """Best-effort encoding detection using optional libraries."""
        # charset-normalizer (ships with many HTTP stacks)
        try:
            from charset_normalizer import detect  # type: ignore[import]

            result = detect(raw)
            enc = result.get("encoding")
            if enc:
                return enc
        except ImportError:
            pass

        # chardet (legacy fallback)
        try:
            from chardet import detect  # type: ignore[import]

            result = detect(raw)
            enc = result.get("encoding")
            if enc:
                return enc
        except ImportError:
            pass

        return None

    @staticmethod
    def _extract_title(content: str) -> str:
        """Derive a title from the first non-empty line of *content*."""
        if not content:
            return ""
        for line in content.splitlines():
            stripped = line.strip()
            if stripped:
                # Remove leading Markdown heading markers.
                cleaned = re.sub(r"^#+\s*", "", stripped)
                return cleaned[:200]
        return ""

    @staticmethod
    def _extract_markdown_title(text: str) -> str:
        """Return the first Markdown heading text, or empty string."""
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                return re.sub(r"^#+\s*", "", stripped).strip()[:200]
        return ""
